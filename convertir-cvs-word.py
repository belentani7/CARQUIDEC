import sys, io, os, re
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from html.parser import HTMLParser
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

GOLD = RGBColor(0xB8, 0x95, 0x3A)
DARK = RGBColor(0x1A, 0x1A, 0x1A)
GRAY = RGBColor(0x66, 0x66, 0x66)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

class TextExtractor(HTMLParser):
    def __init__(self):
        super().__init__()
        self.text = []
        self.in_style = False
        self.in_script = False
        self.current_tag = ''
        self.current_attrs = {}
    def handle_starttag(self, tag, attrs):
        self.current_tag = tag
        self.current_attrs = dict(attrs)
        if tag == 'style': self.in_style = True
        if tag == 'script': self.in_script = True
    def handle_endtag(self, tag):
        if tag == 'style': self.in_style = False
        if tag == 'script': self.in_script = False
    def handle_data(self, data):
        if not self.in_style and not self.in_script:
            t = data.strip()
            if t and len(t) > 1:
                self.text.append(t)

def extract_text(html_file):
    with open(html_file, 'r', encoding='utf-8', errors='ignore') as f:
        html = f.read()
    ext = TextExtractor()
    ext.feed(html)
    return ext.text

def create_docx(texts, output_file, title_text=""):
    doc = Document()
    
    # Set A4 page size
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Segoe UI'
    font.size = Pt(10)
    font.color.rgb = DARK
    
    # Title
    if title_text:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title_text)
        run.bold = True
        run.font.size = Pt(20)
        run.font.color.rgb = DARK
        p.space_after = Pt(4)
    
    # Process text content
    skip_patterns = ['margin:', 'padding:', 'box-sizing:', 'font-family:', 
                     'background', 'color:', 'border:', 'display:', 'position:',
                     'width:', 'height:', 'text-align:', 'text-transform:',
                     'letter-spacing:', 'line-height:', 'overflow:', 'z-index:',
                     'opacity:', 'filter:', 'animation:', 'transition:',
                     'transform:', 'flex:', 'grid:', 'gap:', 'max-width:',
                     'min-height:', 'box-shadow:', 'backdrop-filter:', '{', '}',
                     'px;', 'rem;', 'em;', 'vh;', 'vw;', 'auto;', '@page',
                     ':root', '::-webkit', 'scrollbar', '@keyframes', 'to{',
                     'from{', '0%,', '50%{', '100%{', 'content:', 'inset:',
                     'pointer-events:', 'will-change:', 'cursor:', 'outline:',
                     'resize:', 'vertical-align:', 'white-space:', 'word-',
                     'visibility:', 'clip:', 'float:', 'clear:', 'list-style:',
                     'table-layout:', 'border-collapse:', 'border-spacing:',
                     'caption-side:', 'empty-cells:', 'quotes:', 'counter-',
                     'page-break', 'orphans:', 'widows:', 'marks:', 'size:',
                     'bleed:', 'marks:', 'cssText', 'getPropertyValue']
    
    section_keywords = ['CONTACTO', 'CONTACT', 'PERFIL', 'PROFILE', 'HABILIDADES', 
                       'SKILLS', 'EXPERIENCIA', 'EXPERIENCE', 'EDUCACION', 'EDUCATION',
                       'PROYECTOS', 'PROJECTS', 'PREMIOS', 'AWARDS', 'CERTIFICACIONES',
                       'CERTIFICATIONS', 'IDIOMAS', 'LANGUAGES', 'COMPETENCIAS',
                       'SPECIALTIES', 'LOGROS', 'ACHIEVEMENTS', 'SERVICIOS', 'SERVICES',
                       'METODOLOGIA', 'METHODOLOGY', 'PORTFOLIO', 'SOBRE', 'ABOUT',
                       'HITOS', 'TRAYECTORIA', 'FORMACION', 'ADN', 'ESTANCIAS']
    
    for text in texts:
        # Skip CSS/technical content
        if any(p in text for p in skip_patterns):
            continue
        if len(text) < 3:
            continue
        if re.match(r'^[\d\.]+(px|pt|rem|em|vh|vw|%|s|ms)$', text):
            continue
        if re.match(r'^(#[0-9a-fA-F]{3,8}|rgb|rgba|hsl)', text):
            continue
        
        # Clean text
        text = text.replace('\n', ' ').replace('\t', ' ').strip()
        text = re.sub(r'\s+', ' ', text)
        
        if not text or len(text) < 2:
            continue
            
        p = doc.add_paragraph()
        
        # Check if it's a section header
        is_header = text.upper() in section_keywords or any(text.upper().startswith(k) for k in section_keywords)
        
        if is_header:
            run = p.add_run(text.upper())
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = GOLD
            p.space_before = Pt(8)
            p.space_after = Pt(4)
            # Add bottom border
            pPr = p._p.get_or_add_pPr()
            pBdr = pPr.makeelement(qn('w:pBdr'), {})
            bottom = pBdr.makeelement(qn('w:bottom'), {
                qn('w:val'): 'single',
                qn('w:sz'): '4',
                qn('w:space'): '1',
                qn('w:color'): 'B8953A'
            })
            pBdr.append(bottom)
            pPr.append(pBdr)
        elif text.startswith('+34') or text.startswith('+57') or '@' in text or 'www.' in text or 'carquidec' in text.lower():
            run = p.add_run(text)
            run.font.size = Pt(9)
            run.font.color.rgb = GRAY
        elif re.match(r'^(20\d{2}|19\d{2})', text):
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = GOLD
        elif text in ['Nativo', 'Native', 'Profesional', 'Professional', 'Conversacional', 'Conversational']:
            run = p.add_run(text)
            run.italic = True
            run.font.size = Pt(9)
            run.font.color.rgb = GRAY
        elif any(text.startswith(m) for m in ['Arquitecto', 'Architect', 'Fundador', 'Founder', 'Director', 'Senior', 'paisajista', 'Landscape']):
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = DARK
        elif text.startswith('CARQUIDEC') or text.startswith('Proyectos') or text.startswith('Universidad') or text.startswith('Centro') or text.startswith('Passivhaus') or text.startswith('IAAC'):
            run = p.add_run(text)
            run.font.size = Pt(9)
            run.font.color.rgb = GRAY
        elif re.match(r'^\d+[\+\%]?$', text):
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(14)
            run.font.color.rgb = DARK
        else:
            run = p.add_run(text)
            run.font.size = Pt(10)
            run.font.color.rgb = DARK
        
        p.space_after = Pt(2)
    
    doc.save(output_file)
    print(f'Creado: {output_file} ({os.path.getsize(output_file)//1024}KB)')

# Process all CV files
html_files = {
    'CV-01-DNA-Architecture': r'C:\Users\USER\Downloads\ARCHIVE\deepseek-exports\deepseek_html_20260730_8afb4f.html',
    'CV-02-Architect-CARQUIDEC': r'C:\Users\USER\Downloads\ARCHIVE\deepseek-exports\deepseek_html_20260730_ba34e9.html',
    'CV-03-Architect-CARQUIDEC-v2': r'C:\Users\USER\Downloads\ARCHIVE\deepseek-exports\deepseek_html_20260730_ba34e9 (1).html',
    'CV-04-Portafolio-Arquitectonico': r'C:\Users\USER\Downloads\ARCHIVE\deepseek-exports\deepseek_html_20260730_c1b5b0.html',
    'CV-05-Architect-Navy': r'C:\Users\USER\Downloads\ARCHIVE\deepseek-exports\deepseek_html_20260730_c30a81.html',
    'CV-06-CARQUIDEC-Dark': r'C:\Users\USER\Downloads\ARCHIVE\deepseek-exports\deepseek_html_20260730_dafc01.html',
    'CV-07-CARQUIDEC-Dark-v2': r'C:\Users\USER\Downloads\ARCHIVE\deepseek-exports\deepseek_html_20260730_dafc01 (1).html',
    'CV-08-Architecto-CARQUIDEC': r'C:\Users\USER\Downloads\ARCHIVE\deepseek-exports\deepseek_html_20260730_e966ac.html',
    'CV-09-Curriculum-Original': r'C:\Users\USER\Downloads\mesmo\curriculum.html',
}

out_dir = r'C:\Users\USER\Downloads\mesmo\CV-WORDS'
os.makedirs(out_dir, exist_ok=True)

for name, html_file in html_files.items():
    if os.path.exists(html_file):
        texts = extract_text(html_file)
        out_file = os.path.join(out_dir, f'{name}.docx')
        create_docx(texts, out_file, name.replace('-', ' '))
    else:
        print(f'NO EXISTE: {html_file}')

print('\nListo! Todos los CVs convertidos a Word.')
